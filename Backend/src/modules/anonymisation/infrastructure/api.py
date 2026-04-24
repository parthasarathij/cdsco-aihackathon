import os
import uuid
import logging
from datetime import datetime
import io
from src.utils.logger import get_logger

from fastapi import APIRouter, HTTPException, UploadFile, File
from fastapi.responses import JSONResponse
from fastapi.staticfiles import StaticFiles
from docx import Document as DocxDocument

from ..models import (
    DocumentRequest,
    DocumentResponse,
    AnonymisedChange,
    AnonymisationMode,
    DetectedEntity,
)
from ..detector import EntityDetector
from ..anonymizer import DocumentAnonymiser
from ..mapping_export import export_mapping_to_excel
from ..entity_fields import FINAL_ENTITY_FIELDS, entity_type_category_map
from ..upload_docx_route import router as docx_router, set_singletons

router = APIRouter()
logger = get_logger(__name__)

# Config 
EXCEL_OUTPUT_DIR = "./mapping_exports"
os.makedirs(EXCEL_OUTPUT_DIR, exist_ok=True)

# Singletons to be initialized in server.py
detector:   EntityDetector    = None
anonymiser: DocumentAnonymiser = None

def init_anonymisation():
    """
    Initialize NER model and anonymizer.
    """
    global detector, anonymiser
    logger.info("Loading NER model + adapter...")
    detector  = EntityDetector()
    anonymiser = DocumentAnonymiser()
    set_singletons(detector, anonymiser)
    logger.info("Anonymisation model ready.")

#  Helpers

def _excel_filename(prefix: str = "mapping") -> str:
    ts  = datetime.now().strftime("%Y%m%d_%H%M%S")
    uid = uuid.uuid4().hex[:6].upper()
    return os.path.join(EXCEL_OUTPUT_DIR, f"{prefix}_{ts}_{uid}.xlsx")


def _build_changes(
    entities:       list[DetectedEntity],
    mapping_entries: list,           # MappingEntry list from pseudo pass
    full_anon_map:   dict[str, str], # original_value → full_anon replacement
    mode:            str,
) -> list[AnonymisedChange]:
    """
    Build the clean diff list: one AnonymisedChange per unique original value.
    Counts occurrences in the original entity list.
    """
    # Count how many times each original value appears
    occurrence_count: dict[str, int] = {}
    for e in entities:
        occurrence_count[e.text] = occurrence_count.get(e.text, 0) + 1

    # Build a lookup: original_value → pseudo_token
    pseudo_map: dict[str, str] = {}
    if mapping_entries:
        for entry in mapping_entries:
            pseudo_map[entry.original_value] = entry.token

    changes: list[AnonymisedChange] = []
    seen: set[str] = set()
    serial = 1

    # Keep insertion order from entity list (sorted by start position)
    for e in sorted(entities, key=lambda x: x.start):
        if e.text in seen:
            continue
        seen.add(e.text)

        changes.append(AnonymisedChange(
            serial_no=serial,
            entity_type=e.entity_type,
            original_value=e.text,
            pseudo_value=pseudo_map.get(e.text) if mode in ("pseudo", "both") else None,
            full_anon_value=full_anon_map.get(e.text) if mode in ("full", "both") else None,
            detection_source=e.source,
            confidence=round(e.score, 4),
            occurrences=occurrence_count[e.text],
        ))
        serial += 1

    return changes


def _build_full_anon_map(
    text:     str,
    entities: list[DetectedEntity],
) -> dict[str, str]:
    """
    Run the full-anon pass on each unique entity value to get its
    generalised replacement, without modifying the document text.
    """
    from ..anonymizer import _generaliser_for_entity_type
    result: dict[str, str] = {}
    for e in entities:
        if e.text not in result:
            fn = _generaliser_for_entity_type(e.entity_type)
            result[e.text] = fn(e.text)
    return result


def _process_core(req: DocumentRequest) -> JSONResponse:
    """Shared processing logic used by /process and /upload."""
    if not req.text.strip():
        raise HTTPException(status_code=400, detail="Input text is empty.")

    #  1. Detect
    entities = detector.detect(req.text)

    if not entities:
        resp = DocumentResponse(
            original_text=req.text,
            total_entities_found=0,
            total_values_changed=0,
            changes=[],
            pseudo_document=req.text if req.mode in ("pseudo", "both") else None,
            full_anon_document=req.text if req.mode in ("full", "both") else None,
            message="No PHI/PII entities detected. Document returned unchanged.",
        )
        return JSONResponse(content=resp.dict())

    #  2. Anonymise 
    pseudo_text, full_anon_text, mapping_table = anonymiser.anonymise(
        text=req.text,
        entities=entities,
        mode=req.mode,
        salt=req.salt,
    )

    #  3. Build full-anon lookup (per unique value) 
    full_anon_map = _build_full_anon_map(req.text, entities)

    #  4. Build clean changes list ─
    mapping_entries = mapping_table.entries if mapping_table else []
    changes = _build_changes(entities, mapping_entries, full_anon_map, req.mode)

    #  5. Export Excel ─
    excel_url = None
    if req.mode in ("pseudo", "both") and mapping_table and mapping_table.entries:
        entity_lookup = {e.text: e for e in entities}
        out_path = _excel_filename("mapping")
        export_mapping_to_excel(
            mapping_entries=mapping_table.entries,
            output_path=out_path,
            original_text=req.text,
            meta={"mode": req.mode, "salt": bool(req.salt),
                  "text_length": len(req.text)},
            entity_lookup=entity_lookup,
        )
        excel_url = f"/exports/{os.path.basename(out_path)}"

    #  6. Build response ─
    resp = DocumentResponse(
        original_text=req.text,
        total_entities_found=len(entities),
        total_values_changed=len(changes),

        #  clear diff: what changed into what ─
        changes=changes,

        #  full documents with replacements applied 
        pseudo_document=pseudo_text,
        full_anon_document=full_anon_text,

        #  optional raw detail ─
        entities_detected=entities,
        mapping_table=mapping_table if req.return_mapping else None,

        mapping_excel_url=excel_url,
        message="Processing complete.",
    )
    return JSONResponse(content=resp.dict())


# API Endpoints
@router.get("/health-anon")
def health():
    return {"status": "ok", "model": "dslim/bert-base-NER + adapter"}


@router.get("/entity-fields")
def get_entity_fields():
    """
    Returns the project's final entity fields list (label + entity_type + category).
    Useful for UI dropdowns / validation / reporting.
    """
    return {
        "count": len(FINAL_ENTITY_FIELDS),
        "fields": FINAL_ENTITY_FIELDS,
        "category_by_entity_type": entity_type_category_map(),
    }


@router.get("/exports/list")
def list_exports():
    files = sorted(
        [f for f in os.listdir(EXCEL_OUTPUT_DIR) if f.endswith(".xlsx")],
        reverse=True,
    )
    return {
        "count": len(files),
        "files": [
            {
                "filename": f,
                "url": f"/exports/{f}",
                "size_kb": round(
                    os.path.getsize(os.path.join(EXCEL_OUTPUT_DIR, f)) / 1024, 1
                ),
            }
            for f in files
        ],
    }


@router.post("/process")
def process_document(req: DocumentRequest):
    """Provide anonymisation results including changes, tokenised text, final anonymised text, 
    and mapping file URL."""
    return _process_core(req)


@router.post("/detect-only")
def detect_only(req: DocumentRequest):
    entities = detector.detect(req.text)
    return {
        "text_length":  len(req.text),
        "entity_count": len(entities),
        "entities":     entities,
    }


@router.post("/upload")
async def upload_file(
    file:           UploadFile = File(...),
    mode:           AnonymisationMode = AnonymisationMode.both,
    return_mapping: bool = False,
    salt:           str | None = None,
):
    content  = await file.read()
    filename = (file.filename or "").lower()

    if filename.endswith(".docx"):
        try:
            doc   = DocxDocument(io.BytesIO(content))
            lines = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            lines.append(cell.text)
            text = "\n".join(lines)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to read .docx: {e}")
    elif filename.endswith(".pdf"):
        raise HTTPException(
            status_code=400,
            detail="Use /process-pdf (or /upload-pdf) to anonymise PDFs while preserving structure.",
        )
    else:
        try:
            text = content.decode("utf-8")
        except UnicodeDecodeError:
            text = content.decode("latin-1")

    if not text.strip():
        raise HTTPException(status_code=400, detail="No readable text found.")

    req = DocumentRequest(text=text, mode=mode, return_mapping=return_mapping, salt=salt)
    return _process_core(req)

router.include_router(docx_router)
