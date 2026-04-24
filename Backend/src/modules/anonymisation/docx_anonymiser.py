from __future__ import annotations
from src.utils.logger import get_logger

import io
import logging
import os
import re
import uuid
from dataclasses import dataclass, field
from datetime import datetime
from typing import Any, Optional, Tuple

from docx import Document

from .models import DetectedEntity, MappingEntry, MappingTableResponse
from .detector import EntityDetector, is_token_isolated
from .anonymizer import DocumentAnonymiser, _generaliser_for_entity_type
from .mapping_export import export_mapping_to_excel

logger = get_logger(__name__)


#  Result 

@dataclass
class DocxAnonymisationResult:
    """Generate pseudo-anonymised and fully anonymised DOCX versions
      while returning a combined JSON report of all changes and outputs."""
    pseudo_docx_bytes:    Optional[bytes] = None
    full_anon_docx_bytes: Optional[bytes] = None
    combined_json:        dict            = field(default_factory=dict)


#  Text extraction (reading order must match detector + span mapping) 

def _iter_paragraphs_in_reading_order(doc: Document):
    """Same traversal as legacy extract: body → tables → headers/footers."""
    for para in doc.paragraphs:
        yield para
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para
    for section in doc.sections:
        for container in (
            section.header,            section.first_page_header,
            section.even_page_header,  section.footer,
            section.first_page_footer, section.even_page_footer,
        ):
            if container:
                for para in container.paragraphs:
                    yield para


def _paragraph_plain_text(para) -> str:
    if not para.runs:
        return para.text or ""
    return "".join(r.text for r in para.runs)


def _extract_text_and_spans(doc: Document) -> tuple[str, list[dict[str, Any]]]:
    """Provide text with paragraph offsets and indexes to ensure 
    replacements occur only in the originating paragraph."""
    paras = list(_iter_paragraphs_in_reading_order(doc))
    texts = [_paragraph_plain_text(p) for p in paras]
    spans: list[dict[str, Any]] = []
    pos = 0
    for i, (para, t) in enumerate(zip(paras, texts)):
        if i > 0:
            pos += 1
        start = pos
        pos += len(t)
        spans.append({"index": i, "para": para, "start": start, "end": pos, "text": t})
    full_text = "\n".join(texts)
    return full_text, spans


def _extract_full_text(doc: Document) -> str:
    return _extract_text_and_spans(doc)[0]


def _paragraphs_list(doc: Document) -> list:
    return list(_iter_paragraphs_in_reading_order(doc))


def _group_entities_by_paragraph(
    entities: list[DetectedEntity],
    spans: list[dict[str, Any]],
) -> dict[int, list[DetectedEntity]]:
    """Map paragraph index → entities fully contained in that paragraph's global span."""
    by_index: dict[int, list[DetectedEntity]] = {}
    for sp in spans:
        idx = sp["index"]
        s0, s1 = sp["start"], sp["end"]
        local: list[DetectedEntity] = []
        for e in entities:
            if e.start >= s0 and e.end <= s1:
                local.append(
                    e.model_copy(update={
                        "start": e.start - s0,
                        "end":   e.end - s0,
                    })
                )
        if local:
            local.sort(key=lambda x: x.start)
            by_index[idx] = local
    return by_index


def _distinct_row_cells(row) -> list:
    """De-duplicate cells when Word returns the same merged cell multiple times."""
    seen: set = set()
    out: list = []
    for cell in row.cells:
        el = getattr(cell, "_element", None) or getattr(cell, "_tc", None) or id(cell)
        if el in seen:
            continue
        seen.add(el)
        out.append(cell)
    return out


def _cell_joined_text(cell) -> str:
    parts = [_paragraph_plain_text(p) for p in cell.paragraphs]
    return "\n".join(parts).strip()


def _cell_offset_to_para_and_local(cell, char_offset: int):
    pos = 0
    paras = list(cell.paragraphs)
    for i, para in enumerate(paras):
        pt = _paragraph_plain_text(para)
        plen = len(pt)
        if plen == 0:
            if i < len(paras) - 1:
                pos += 1
            continue
        if pos <= char_offset < pos + plen:
            return para, char_offset - pos
        pos += plen
        if i < len(paras) - 1:
            pos += 1
    return None, -1


def _find_isolated_in_cell(cell, needle: str):
    cn = _cell_joined_text(cell)
    if not needle or needle not in cn:
        return None
    p = 0
    while True:
        i = cn.find(needle, p)
        if i < 0:
            return None
        if is_token_isolated(cn, i, i + len(needle)):
            para, loc = _cell_offset_to_para_and_local(cell, i)
            if para is not None and loc >= 0:
                return para, loc, len(needle)
        p = i + 1
    return None


def _global_span_for_para_offset(
    para_spans: list[dict[str, Any]],
    para,
    local_start: int,
    length: int,
) -> Optional[Tuple[int, int]]:
    for sp in para_spans:
        if sp["para"] is para:
            g0 = sp["start"] + local_start
            return g0, g0 + length
    return None


def _augment_entities_from_table_rows(
    detector: EntityDetector,
    doc: Document,
    para_spans: list[dict[str, Any]],
    entities: list[DetectedEntity],
) -> list[DetectedEntity]:
    """
    Re-run detection on each table row with cells joined by tab so patterns like
    'Death Date ... \\t 15 June 2023' match when label and value are in
    adjacent cells.
    """
    extra: list[DetectedEntity] = []
    occupied: list[tuple[int, int]] = [(e.start, e.end) for e in entities]

    def _overlaps(a0: int, a1: int) -> bool:
        for b0, b1 in occupied:
            if a0 < b1 and b0 < a1:
                return True
        return False

    for table in doc.tables:
        for row in table.rows:
            cells = _distinct_row_cells(row)
            if not cells:
                continue
            cell_texts = [_cell_joined_text(c) for c in cells]
            joined = "\t".join(cell_texts)
            if len(joined.strip()) < 2:
                continue
            for e in detector.detect(joined):
                if not is_token_isolated(joined, e.start, e.end):
                    continue
                col = joined[: e.start].count("\t")
                if col < 0 or col >= len(cells):
                    continue
                hit = _find_isolated_in_cell(cells[col], e.text)
                if hit is None:
                    continue
                para, loc, L = hit
                gs_ge = _global_span_for_para_offset(para_spans, para, loc, L)
                if gs_ge is None:
                    continue
                g0, g1 = gs_ge
                if _overlaps(g0, g1):
                    continue
                occupied.append((g0, g1))
                extra.append(DetectedEntity(
                    text=e.text,
                    entity_type=e.entity_type,
                    start=g0,
                    end=g1,
                    score=e.score,
                    source=e.source,
                ))
    if not extra:
        return entities
    return sorted(entities + extra, key=lambda x: x.start)


#  Core in-place replacement 

def _replace_in_para_by_spans(para, operations: list[tuple[int, int, str]]) -> None:
    """
    Replace using exact [start:end) offsets from detection (paragraph-local).
    Avoids all substring/regex false positives when the same text appears in
    both a label and a value, and fixes glued tokens (e.g. DepressionENT_…).
    """
    if not para.runs or not operations:
        return
    full = "".join(r.text for r in para.runs)
    result = full
    for start, end, token in sorted(operations, key=lambda x: -x[0]):
        if start < 0 or start >= end or not token:
            continue
        if end > len(result):
            continue
        if result[start:end] != full[start:end]:
            continue
        result = result[:start] + token + result[end:]
    if result == full:
        return
    para.runs[0].text = result
    for r in para.runs[1:]:
        r.text = ""


def _replace_in_para(para, replacements: dict[str, str]) -> None:
    """
    Legacy: dict replace with word-boundary guards. Prefer
    `_replace_in_para_by_spans` when entity offsets are known.
    """
    if not para.runs:
        return

    full = "".join(r.text for r in para.runs)
    replaced = full

    for original, token in sorted(replacements.items(), key=lambda x: -len(x[0])):
        if not original:
            continue
        esc = re.escape(original)
        pattern = rf"(?<!\w){esc}(?!\w)"
        try:
            replaced = re.sub(pattern, token, replaced)
        except re.error:
            replaced = replaced.replace(original, token)

    if replaced == full:
        return  # nothing changed — leave runs untouched

    para.runs[0].text = replaced
    for r in para.runs[1:]:
        r.text = ""


def _apply_replacements_to_doc(doc: Document, replacements: dict[str, str]) -> None:
    """
    Walk every text-bearing container in the document and apply replacements.
    Covers: body paragraphs, table cells, headers, footers.
    """
    for para in doc.paragraphs:
        _replace_in_para(para, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_para(para, replacements)

    for section in doc.sections:
        for container in (
            section.header,            section.first_page_header,
            section.even_page_header,  section.footer,
            section.first_page_footer, section.even_page_footer,
        ):
            if container:
                for para in container.paragraphs:
                    _replace_in_para(para, replacements)


def _doc_to_bytes(doc: Document) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


#  JSON builder ─

def _build_json(
    original_text:  str,
    entities:       list,
    pseudo_text:    Optional[str],
    full_anon_text: Optional[str],
    mapping_table:  Optional[MappingTableResponse],
    mode:           str,
    excel_url:      Optional[str],
    return_mapping: bool,
) -> dict:
    occ: dict[str, int] = {}
    for e in entities:
        occ[e.text] = occ.get(e.text, 0) + 1

    pseudo_map: dict[str, str] = {}
    if mapping_table:
        for entry in mapping_table.entries:
            pseudo_map[entry.original_value] = entry.token

    full_anon_map: dict[str, str] = {}
    for e in entities:
        if e.text not in full_anon_map:
            fn = _generaliser_for_entity_type(e.entity_type)
            full_anon_map[e.text] = fn(e.text)

    changes: list[dict] = []
    seen: set[str] = set()
    serial = 1
    for e in sorted(entities, key=lambda x: x.start):
        if e.text in seen:
            continue
        seen.add(e.text)
        changes.append({
            "serial_no":        serial,
            "entity_type":      e.entity_type,
            "original_value":   e.text,
            "pseudo_value":     pseudo_map.get(e.text),
            "full_anon_value":  full_anon_map.get(e.text),
            "detection_source": e.source,
            "confidence":       round(e.score, 4),
            "occurrences":      occ[e.text],
        })
        serial += 1

    payload: dict = {
        "original_text":        original_text,
        "total_entities_found": len(entities),
        "total_values_changed": len(changes),
        "changes":              changes,
        "pseudo_document":      pseudo_text,
        "full_anon_document":   full_anon_text,
        "mapping_excel_url":    excel_url,
        "message":              "Processing complete." if entities
                                else "No PHI/PII entities detected.",
    }

    if return_mapping and mapping_table:
        payload["mapping_table"] = {
            "entries": [
                {
                    "token":          e.token,
                    "original_value": e.original_value,
                    "entity_type":    e.entity_type,
                    "source":         getattr(e, "source", None),
                    "score":          getattr(e, "score",  None),
                }
                for e in mapping_table.entries
            ]
        }

    return payload


#  Main entry point 

def process_docx_bytes(
    file_bytes:       bytes,
    mode:             str           = "both",
    salt:             Optional[str] = None,
    return_mapping:   bool          = True,
    detector=None,
    anonymiser=None,
    excel_output_dir: str           = "./mapping_exports",
) -> DocxAnonymisationResult:
    """
    Full pipeline:

      1. Extract plain text from the .docx.
      2. Run the hybrid detector (regex + NER) to find PII/PHI spans.
      3. Run the anonymiser to get pseudo tokens + full-anon replacements.
      4. Apply each replacement dict to a FRESH COPY of the original Document.
           → Both output docs have the EXACT SAME layout as the original.
           → Only PII text values are swapped in-place.
           → No extra sections, banners, or page breaks are added.
      5. Return DocxAnonymisationResult with both byte streams + JSON report.

    Parameters
    ----------
    file_bytes        Raw bytes of the original .docx file.
    mode              "pseudo" | "full" | "both"
    salt              Optional HMAC salt for deterministic pseudo tokens.
    return_mapping    Include the full mapping table in the JSON report.
    detector          Pre-loaded EntityDetector (reuse across requests).
    anonymiser        Ignored for .docx (a fresh DocumentAnonymiser is used per
                        file so tokens do not leak across uploads).
    excel_output_dir  Directory where Excel mapping exports are written.
    """

    if detector is None:
        detector = EntityDetector()

    #  Load doc, extract text + paragraph spans (for scoped .docx replace) 
    doc_original = Document(io.BytesIO(file_bytes))
    full_text, para_spans = _extract_text_and_spans(doc_original)

    if not full_text.strip():
        return DocxAnonymisationResult(
            combined_json={
                "message":              "No readable text found.",
                "total_entities_found": 0,
                "total_values_changed": 0,
                "changes":              [],
                "pseudo_document":      None,
                "full_anon_document":   None,
            }
        )

    #  Detect on full text, then add hits from table rows (label\\tvalue) ─
    entities = detector.detect(full_text)
    entities = _augment_entities_from_table_rows(
        detector, doc_original, para_spans, entities,
    )

    if not entities:
        no_pii_bytes = _doc_to_bytes(doc_original)
        return DocxAnonymisationResult(
            pseudo_docx_bytes    = no_pii_bytes if mode in ("pseudo", "both") else None,
            full_anon_docx_bytes = no_pii_bytes if mode in ("full",   "both") else None,
            combined_json={
                "message":              "No PHI/PII entities detected. Document returned unchanged.",
                "total_entities_found": 0,
                "total_values_changed": 0,
                "changes":              [],
                "pseudo_document":      full_text,
                "full_anon_document":   full_text,
            },
        )

    # Fresh anonymiser per document so vault/tokens are not shared across uploads.
    anon = DocumentAnonymiser()

    by_index = _group_entities_by_paragraph(entities, para_spans)

    merged_entries: list[MappingEntry] = []
    seen_original: set[str] = set()
    pseudo_blocks:  list[str] = []
    full_blocks:    list[str] = []

    for sp in para_spans:
        t = sp["text"]
        idx = sp["index"]
        local = by_index.get(idx, [])

        if not local:
            pseudo_blocks.append(t)
            full_blocks.append(t)
            continue

        pseudo_t, full_t, mapping = anon.anonymise(t, local, mode=mode, salt=salt)
        pseudo_blocks.append(pseudo_t if pseudo_t is not None else t)
        full_blocks.append(full_t if full_t is not None else t)

        if mapping and mapping.entries:
            for entry in mapping.entries:
                if entry.original_value in seen_original:
                    continue
                seen_original.add(entry.original_value)
                merged_entries.append(entry)

    mapping_table = (
        MappingTableResponse(entries=merged_entries) if merged_entries else None
    )
    pseudo_text    = "\n".join(pseudo_blocks)
    full_anon_text = "\n".join(full_blocks)

    pseudo_replacements: dict[str, str] = {}
    if mode in ("pseudo", "both") and mapping_table:
        for entry in mapping_table.entries:
            if entry.original_value and entry.token:
                pseudo_replacements[entry.original_value] = entry.token

    full_anon_replacements: dict[str, str] = {}
    if mode in ("full", "both"):
        for e in entities:
            if e.text and e.text not in full_anon_replacements:
                fn = _generaliser_for_entity_type(e.entity_type)
                full_anon_replacements[e.text] = fn(e.text)

    pseudo_docx_bytes:    Optional[bytes] = None
    full_anon_docx_bytes: Optional[bytes] = None

    if mode in ("pseudo", "both") and pseudo_replacements:
        doc_pseudo = Document(io.BytesIO(file_bytes))
        paras_p = _paragraphs_list(doc_pseudo)
        for sp in para_spans:
            local = by_index.get(sp["index"], [])
            if not local:
                continue
            t = sp["text"]
            ops: list[tuple[int, int, str]] = []
            for ent in local:
                tok = pseudo_replacements.get(ent.text)
                if not tok:
                    continue
                if ent.start < 0 or ent.end > len(t) or ent.start >= ent.end:
                    continue
                if t[ent.start:ent.end] != ent.text:
                    logger.debug(
                        "Skipping pseudo replace: span text mismatch idx=%s",
                        sp["index"],
                    )
                    continue
                ops.append((ent.start, ent.end, tok))
            if ops:
                _replace_in_para_by_spans(paras_p[sp["index"]], ops)

        pseudo_docx_bytes = _doc_to_bytes(doc_pseudo)

    if mode in ("full", "both") and full_anon_replacements:
        doc_full = Document(io.BytesIO(file_bytes))
        paras_f = _paragraphs_list(doc_full)
        for sp in para_spans:
            local = by_index.get(sp["index"], [])
            if not local:
                continue
            t = sp["text"]
            ops = []
            for ent in local:
                repl_tok = full_anon_replacements.get(ent.text)
                if repl_tok is None:
                    continue
                if ent.start < 0 or ent.end > len(t) or ent.start >= ent.end:
                    continue
                if t[ent.start:ent.end] != ent.text:
                    continue
                ops.append((ent.start, ent.end, repl_tok))
            if ops:
                _replace_in_para_by_spans(paras_f[sp["index"]], ops)

        full_anon_docx_bytes = _doc_to_bytes(doc_full)

    #  Export Excel mapping ─
    excel_url: Optional[str] = None
    if mode in ("pseudo", "both") and mapping_table and mapping_table.entries:
        os.makedirs(excel_output_dir, exist_ok=True)
        ts       = datetime.now().strftime("%Y%m%d_%H%M%S")
        uid      = uuid.uuid4().hex[:6].upper()
        out_path = os.path.join(excel_output_dir, f"mapping_{ts}_{uid}.xlsx")
        export_mapping_to_excel(
            mapping_entries = mapping_table.entries,
            output_path     = out_path,
            original_text   = full_text,
            meta            = {"mode": mode, "salt": bool(salt), "text_length": len(full_text)},
            entity_lookup   = {e.text: e for e in entities},
        )
        excel_url = f"/exports/{os.path.basename(out_path)}"

    #  Build JSON report 
    combined_json = _build_json(
        original_text  = full_text,
        entities       = entities,
        pseudo_text    = pseudo_text,
        full_anon_text = full_anon_text,
        mapping_table  = mapping_table,
        mode           = mode,
        excel_url      = excel_url,
        return_mapping = return_mapping,
    )

    return DocxAnonymisationResult(
        pseudo_docx_bytes    = pseudo_docx_bytes,
        full_anon_docx_bytes = full_anon_docx_bytes,
        combined_json        = combined_json,
    )