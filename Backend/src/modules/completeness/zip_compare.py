from __future__ import annotations
from src.utils.logger import get_logger
logger = get_logger(__name__)

import zipfile
from io import BytesIO
from pathlib import PurePosixPath
from typing import Iterable
from xml.etree import ElementTree

from docx import Document
from pypdf import PdfReader

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _is_safe_zip_member(name: str) -> bool:
    if not name or name.startswith("/"):
        return False
    parts = PurePosixPath(name).parts
    if ".." in parts:
        return False
    return True


def _zip_file_members(names: Iterable[str]) -> list[str]:
    out: list[str] = []
    for raw in names:
        name = raw.replace("\\", "/").strip()
        if not name or name.endswith("/"):
            continue
        if "__MACOSX/" in name or name.startswith("__MACOSX"):
            continue
        if not _is_safe_zip_member(name):
            continue
        out.append(name)
    return out


def basename_key(zip_internal_path: str) -> str:
    return PurePosixPath(zip_internal_path).name.lower()


def list_members_only_by_basename_diff(
    zip_a_path: str,
    zip_b_path: str,
) -> tuple[list[str], list[str]]:
    """
    Compare two zips by file name only (case-insensitive basename).

    Returns (names_only_in_a, names_only_in_b) as full member paths inside each zip.
    If multiple members share the same basename in one zip, each path is listed if that basename is missing from the other zip.
    """
    with zipfile.ZipFile(zip_a_path, "r") as za, zipfile.ZipFile(zip_b_path, "r") as zb:
        members_a = _zip_file_members(za.namelist())
        members_b = _zip_file_members(zb.namelist())

    keys_b = {basename_key(m) for m in members_b}
    keys_a = {basename_key(m) for m in members_a}

    only_a = [m for m in members_a if basename_key(m) not in keys_b]
    only_b = [m for m in members_b if basename_key(m) not in keys_a]
    return only_a, only_b


def extract_pdf_first_n_pages_text(data: bytes, *, max_pages: int = 5, max_chars: int = 14_000) -> str:
    reader = PdfReader(BytesIO(data))
    chunks: list[str] = []
    n = min(max_pages, len(reader.pages))
    for i in range(n):
        try:
            t = reader.pages[i].extract_text() or ""
        except Exception:
            t = ""
        chunks.append(t)
    text = "\n".join(chunks).strip()
    return text[:max_chars]


def _docx_xml_wt_text(data: bytes, *, max_chars: int) -> str:
    """Pull visible text from word/document.xml (covers many table / SDT / textbox cases)."""
    try:
        with zipfile.ZipFile(BytesIO(data), "r") as zf:
            xml_bytes = zf.read("word/document.xml")
    except (KeyError, OSError, zipfile.BadZipFile):
        return ""
    try:
        root = ElementTree.fromstring(xml_bytes)
    except ElementTree.ParseError:
        return ""
    chunks: list[str] = []
    total = 0
    tag_t = f"{{{_W_NS}}}t"
    for el in root.iter(tag_t):
        for piece in (el.text, el.tail):
            if not piece:
                continue
            chunks.append(piece)
            total += len(piece)
            if total >= max_chars:
                return "".join(chunks)[:max_chars].strip()
    return "".join(chunks)[:max_chars].strip()


def extract_docx_substantive_text(data: bytes, *, max_chars: int = 14_000) -> str:
    """
    Body paragraphs, table cell text (including nested tables), then header/footer blocks.
    If that yields little text (common for heavily formatted clinical summaries), fall back
    to raw ``w:t`` runs from ``document.xml``.
    """
    doc = Document(BytesIO(data))
    lines: list[str] = []

    def emit(raw: str) -> None:
        s = (raw or "").strip()
        if s:
            lines.append(s)

    for p in doc.paragraphs:
        emit(p.text)

    def walk_table(table) -> None:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    emit(p.text)
                for nested in cell.tables:
                    walk_table(nested)

    for table in doc.tables:
        walk_table(table)

    try:
        for sec in doc.sections:
            for part in (sec.header, sec.footer):
                try:
                    for p in part.paragraphs:
                        emit(p.text)
                    for table in part.tables:
                        walk_table(table)
                except Exception:
                    continue
    except Exception:
        pass

    structured = "\n".join(lines).strip()
    structured = structured[:max_chars]
    if len(structured) < 80:
        raw = _docx_xml_wt_text(data, max_chars=max_chars)
        if len(raw) > len(structured):
            return raw[:max_chars]
    return structured[:max_chars] if structured else _docx_xml_wt_text(data, max_chars=max_chars)[:max_chars]


def snippet_for_llm_from_zip_member(zip_path: str, member: str) -> tuple[str | None, str | None]:
    """
    Returns (snippet_text, error_or_skip_reason).

    snippet_text is None when the member is not PDF/DOCX or cannot be read.
    """
    suffix = PurePosixPath(member).suffix.lower()
    with zipfile.ZipFile(zip_path, "r") as z:
        try:
            data = z.read(member)
        except (KeyError, RuntimeError, zipfile.BadZipFile) as e:
            return None, f"read_failed: {e!s}"

    if suffix == ".pdf":
        try:
            return extract_pdf_first_n_pages_text(data, max_pages=5), None
        except Exception as e:
            return None, f"pdf_extract_failed: {e!s}"

    if suffix == ".docx":
        try:
            return extract_docx_substantive_text(data), None
        except Exception as e:
            return None, f"docx_extract_failed: {e!s}"

    return None, "unsupported_type"


def compare_zips(zip_a_bytes: bytes, zip_b_bytes: bytes, zip_a_name: str = "zip_a", zip_b_name: str = "zip_b") -> dict:
    """
    Compare two ZIP files and return differences.
    Returns dict with only_in_zip_a and only_in_zip_b lists.
    """
    try:
        with zipfile.ZipFile(BytesIO(zip_a_bytes), "r") as za:
            members_a = _zip_file_members(za.namelist())
        with zipfile.ZipFile(BytesIO(zip_b_bytes), "r") as zb:
            members_b = _zip_file_members(zb.namelist())
    except zipfile.BadZipFile as e:
        raise ValueError(f"Invalid ZIP file: {e}")

    keys_b = {basename_key(m) for m in members_b}
    keys_a = {basename_key(m) for m in members_a}

    only_a = [m for m in members_a if basename_key(m) not in keys_b]
    only_b = [m for m in members_b if basename_key(m) not in keys_a]

    # Convert to the expected format
    result_a = []
    result_b = []

    for path in only_a:
        result_a.append({
            "document_added": zip_a_name,
            "path_in_zip": path,
            "Description": f"File only in {zip_a_name}"
        })

    for path in only_b:
        result_b.append({
            "document_added": zip_b_name,
            "path_in_zip": path,
            "Description": f"File only in {zip_b_name}"
        })

    return {
        "only_in_zip_a": result_a,
        "only_in_zip_b": result_b
    }
