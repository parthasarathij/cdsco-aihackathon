from __future__ import annotations
from src.utils.logger import get_logger
logger = get_logger(__name__)

import zipfile
from contextlib import suppress
from io import BytesIO
from pathlib import Path, PurePosixPath

from docx import Document
from pypdf import PdfReader

# Limits to keep CPU/memory bounded on large dossiers
MAX_PDF_PAGES = 80
MAX_TEXT_CHARS = 450_000


def extract_text_from_pdf_bytes(data: bytes) -> str:
    reader = PdfReader(BytesIO(data))
    n = min(len(reader.pages), MAX_PDF_PAGES)
    parts: list[str] = []
    total = 0
    for i in range(n):
        try:
            t = reader.pages[i].extract_text() or ""
        except Exception:
            t = ""
        parts.append(t)
        total += len(t)
        if total >= MAX_TEXT_CHARS:
            break
    return "\n".join(parts)[:MAX_TEXT_CHARS]


def _walk_docx_tables(table, emit) -> None:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                emit(p.text or "")
            for nested in cell.tables:
                _walk_docx_tables(nested, emit)


def extract_text_from_docx_bytes(data: bytes) -> str:
    doc = Document(BytesIO(data))
    lines: list[str] = []

    def emit(raw: str) -> None:
        s = (raw or "").strip()
        if s:
            lines.append(s)

    for p in doc.paragraphs:
        emit(p.text)
    for table in doc.tables:
        _walk_docx_tables(table, emit)
    with suppress(Exception):
        for sec in doc.sections:
            for part in (sec.header, sec.footer):
                with suppress(Exception):
                    for p in part.paragraphs:
                        emit(p.text)
                    for table in part.tables:
                        _walk_docx_tables(table, emit)
    text = "\n".join(lines)
    if len(text.strip()) < 200:
        text = _docx_raw_wt_text(data) or text
    return text[:MAX_TEXT_CHARS]


def _docx_raw_wt_text(data: bytes) -> str:
    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    try:
        with zipfile.ZipFile(BytesIO(data), "r") as zf:
            xml_bytes = zf.read("word/document.xml")
    except (KeyError, OSError, zipfile.BadZipFile):
        return ""
    from xml.etree import ElementTree

    try:
        root = ElementTree.fromstring(xml_bytes)
    except ElementTree.ParseError:
        return ""
    chunks: list[str] = []
    tag_t = f"{{{W_NS}}}t"
    for el in root.iter(tag_t):
        chunks.extend(piece for piece in (el.text, el.tail) if piece)
    return "".join(chunks)[:MAX_TEXT_CHARS]


def extract_text_from_upload(filename: str, data: bytes) -> str:
    ext = PurePosixPath(filename).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf_bytes(data)
    if ext == ".docx":
        return extract_text_from_docx_bytes(data)
    raise ValueError(f"Unsupported type {ext}; use PDF or DOCX.")


def extract_text_from_path(path: str | Path) -> str:
    p = Path(path)
    return extract_text_from_upload(p.name, p.read_bytes())
